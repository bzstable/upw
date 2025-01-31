import unittest
import json
import os
from docx import Document
from json_to_word import DocumentGenerator

class TestDocumentGenerator(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # Load sample data
        with open('sample.json', 'r') as f:
            cls.json_data = json.load(f)
        
        # Generate document
        cls.generator = DocumentGenerator(cls.json_data)
        cls.generator.generate_document('test_output.docx')
        
        # Load generated document
        cls.doc = Document('test_output.docx')
        cls.paragraphs = [p.text for p in cls.doc.paragraphs]
        cls.tables = cls.doc.tables

    def get_table_data(self, table):
        """Extract data from a table."""
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        return data

    def test_type_a_content(self):
        """Test if Type A content is correctly included."""
        type_a_data = self.json_data['objects']['TypeA']
        
        # Check if device IDs are present
        table_data = self.get_table_data(self.tables[0])  # First table should be Type A
        for item in type_a_data:
            device_id = item['id']
            self.assertTrue(
                any(device_id in row for row in table_data),
                f"Device ID {device_id} not found in document"
            )

    def test_type_b_content(self):
        """Test if Type B content is correctly included."""
        type_b_data = self.json_data['objects']['TypeB']
        
        for item in type_b_data:
            # Check if component ID is present
            component_id = item['component_id']
            found = False
            for table in self.tables:
                table_data = self.get_table_data(table)
                if any(component_id in row for row in table_data):
                    found = True
                    break
            self.assertTrue(found, f"Component ID {component_id} not found in document")

    def test_type_c_content(self):
        """Test if Type C content is correctly included."""
        type_c_data = self.json_data['objects']['TypeC']
        
        for item in type_c_data:
            # Check if transaction ID is present
            transaction_id = item['transaction_id']
            found = False
            for table in self.tables:
                table_data = self.get_table_data(table)
                if any(transaction_id in row for row in table_data):
                    found = True
                    break
            self.assertTrue(found, f"Transaction ID {transaction_id} not found in document")

    def test_type_d_content(self):
        """Test if Type D content is correctly included."""
        type_d_data = self.json_data['objects']['TypeD']
        
        for item in type_d_data:
            # Check if test ID is present
            test_id = item['test_id']
            found = False
            for table in self.tables:
                table_data = self.get_table_data(table)
                if any(test_id in row for row in table_data):
                    found = True
                    break
            self.assertTrue(found, f"Test ID {test_id} not found in document")

    def test_type_e_content(self):
        """Test if Type E content is correctly included."""
        type_e_data = self.json_data['objects']['TypeE']
        
        for item in type_e_data:
            # Check if employee ID is present
            employee_id = item['employee_id']
            found = False
            for table in self.tables:
                table_data = self.get_table_data(table)
                if any(employee_id in row for row in table_data):
                    found = True
                    break
            self.assertTrue(found, f"Employee ID {employee_id} not found in document")

    @classmethod
    def tearDownClass(cls):
        # Clean up test file
        if os.path.exists('test_output.docx'):
            os.remove('test_output.docx')

if __name__ == '__main__':
    unittest.main() 