import json
import yaml
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import os

class DocumentGenerator:
    def __init__(self, json_data: Dict, template_file: str = 'text_templates.yaml'):
        """Initialize the document generator with JSON data and templates."""
        self.json_data = json_data
        with open(template_file, 'r') as f:
            self.templates = yaml.safe_load(f)
        self.doc = Document()

    def add_title(self, text: str) -> None:
        """Add a title to the document with specific formatting."""
        title = self.doc.add_heading(text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()

    def add_table(self, headers: List[str], rows: List[List[str]], title: str = None) -> None:
        """Add a table to the document with headers and rows."""
        if title:
            title_para = self.doc.add_paragraph()
            title_para.add_run(title).bold = True
        
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

        self.doc.add_paragraph()

    def process_type_a(self, data: Dict) -> None:
        """Process Type A objects."""
        template = self.templates['type_a']
        self.add_title(template['title'])
        
        # Table 1: Device Information
        rows = [[item['id'], item['name'], item['description']] for item in data]
        self.add_table(template['table1_headers'], rows, template['table1_title'])
        
        # Table 2: Parameters
        for item in data:
            params = item['parameters']
            rows = [[params['status'], params['priority'], params['last_updated']]]
            self.add_table(template['table2_headers'], rows, template['table2_title'])
        
        # Table 3: Additional Info
        rows = [["Total Devices", len(data)]]
        self.add_table(template['table3_headers'], rows, template['table3_title'])

    def process_type_b(self, data: Dict) -> None:
        """Process Type B objects."""
        template = self.templates['type_b']
        self.add_title(template['title'])
        
        for item in data:
            # Table 1: Component Details
            rows = [[item['component_id'], item['installation_date']]]
            self.add_table(template['table1_headers'], rows, template['table1_title'])
            
            # Table 2: Technical Specs
            specs = item['specs']
            rows = [[k, v] for k, v in specs.items()]
            self.add_table(template['table2_headers'], rows, template['table2_title'])
            
            # Table 3: Manufacturer Info
            rows = [["Manufacturer", specs['manufacturer']]]
            self.add_table(template['table3_headers'], rows, template['table3_title'])

    def process_type_c(self, data: Dict) -> None:
        """Process Type C objects."""
        template = self.templates['type_c']
        self.add_title(template['title'])
        
        for item in data:
            # Table 1: Transaction Overview
            rows = [[item['transaction_id'], item['amount'], item['currency']]]
            self.add_table(template['table1_headers'], rows, template['table1_title'])
            
            # Table 2: Party Information
            rows = [[i+1, party] for i, party in enumerate(item['parties'])]
            self.add_table(template['table2_headers'], rows, template['table2_title'])
            
            # Table 3: Status
            rows = [["Approved", "Yes" if item['approved'] else "No"]]
            self.add_table(template['table3_headers'], rows, template['table3_title'])

    def process_type_d(self, data: Dict) -> None:
        """Process Type D objects."""
        template = self.templates['type_d']
        self.add_title(template['title'])
        
        for item in data:
            # Table 1: Test Information
            rows = [[item['test_id'], item['environment']]]
            self.add_table(template['table1_headers'], rows, template['table1_title'])
            
            # Table 2: Metrics
            metrics = item['metrics']
            rows = [[k.replace('_', ' ').title(), v] for k, v in metrics.items()]
            self.add_table(template['table2_headers'], rows, template['table2_title'])
            
            # Table 3: Additional Details
            rows = [["Environment Type", item['environment'].title()]]
            self.add_table(template['table3_headers'], rows, template['table3_title'])

    def process_type_e(self, data: Dict) -> None:
        """Process Type E objects."""
        template = self.templates['type_e']
        self.add_title(template['title'])
        
        for item in data:
            # Table 1: Employee Details
            rows = [[item['employee_id'], item['department'], item['role']]]
            self.add_table(template['table1_headers'], rows, template['table1_title'])
            
            # Table 2: Project Information
            rows = [[i+1, project] for i, project in enumerate(item['projects'])]
            self.add_table(template['table2_headers'], rows, template['table2_title'])
            
            # Table 3: Additional Information
            rows = [["Total Projects", len(item['projects'])]]
            self.add_table(template['table3_headers'], rows, template['table3_title'])

    def generate_document(self, output_file: str = 'output.docx') -> None:
        """Generate the complete Word document."""
        objects = self.json_data.get('objects', {})
        
        # Process each type of object
        if 'TypeA' in objects:
            self.process_type_a(objects['TypeA'])
        if 'TypeB' in objects:
            self.process_type_b(objects['TypeB'])
        if 'TypeC' in objects:
            self.process_type_c(objects['TypeC'])
        if 'TypeD' in objects:
            self.process_type_d(objects['TypeD'])
        if 'TypeE' in objects:
            self.process_type_e(objects['TypeE'])
        
        # Save the document
        self.doc.save(output_file)


if __name__ == '__main__':
    # Read JSON data
    with open('sample.json', 'r') as f:
        json_data = json.load(f)
    
    # Specify output path (you can change this to any path you want)
    output_path = 'generated_document.docx'
    
    # Generate document
    generator = DocumentGenerator(json_data)
    generator.generate_document(output_path)
    
    # Print confirmation message with absolute path
    abs_path = os.path.abspath(output_path)
    print(f"\nDocument generated successfully!")
    print(f"Location: {abs_path}") 